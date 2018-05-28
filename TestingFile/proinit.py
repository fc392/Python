import os
from docx import Document
from docx.oxml.ns import qn

with open('201802.txt','rb') as f:
    txt = f.read().decode('utf-8')
f.close()
path_txt = os.getcwd()+ '\\Test_txt\\'
path_docx = os.getcwd()+ '\\Test_docx\\'
def main_txt(i):
    path1 = path_txt + str(i) + '.txt'#需要写入文件的名字
    with open(path1,'w+',encoding='utf-8') as fw:
        fw.write(txt[i*800:(i+1)*800])
    fw.close()
    pass

def main_docx(i):
    path2 = path_docx + str(i) + '.docx'
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()
    doc.styles['Normal'].font.name = u'宋体'
    #doc.styles['Normal']._element.rPr.rFonts.set(qn('w:wastAsia'), '宋体')
    t = txt[i*10000:(i+1)*10000]
    doc.add_paragraph(t,style='Normal')
    doc.save(path2)


if __name__ == '__main__':
    n = len(txt)//10000
    for i in range(n+1):
        main_docx(i)